import polars as pl
from docx import Document
from docx2pdf import convert as docx2pdf_convert
import subprocess
import shutil
import os

# А А Г А Б

def create_excel():
    df = pl.DataFrame({
        "Имя": ["Яблоки", "Бананы", "Молоко", "Хлеб", "Сыр", "Яйца", "Масло", "Кофе", "Чай", "Шоколад"],
        "Количество": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10],
        "Цена": [10, 20, 30, 40, 50, 60, 70, 80, 90, 100]
    })

    df = df.with_columns((pl.col("Количество") * pl.col("Цена")).alias("Стоимость"))

    df.write_excel('excel.xlsx')

    return df


def create_word(df):
    doc = Document()

    doc.add_heading('Исходные данные', level=1)

    columns = df.columns
    table = doc.add_table(rows=1, cols=len(columns))
    header_cells = table.rows[0].cells
    for idx, col_name in enumerate(columns):
        header_cells[idx].text = str(col_name)

    for row in df.iter_rows(named=True):
        row_cells = table.add_row().cells
        for idx, col_name in enumerate(columns):
            row_cells[idx].text = str(row[col_name])

    doc.add_heading('Итоги', level=1)

    total_cost = df.select(pl.col("Стоимость").sum()).item()
    doc.add_paragraph(f"Итоговая сумма (Стоимость): {total_cost}")

    max_price = df.select(pl.col("Цена").max()).item()
    most_expensive = df.filter(pl.col("Цена") == max_price).row(0, named=True)
    doc.add_paragraph(
        f"Самая дорогая позиция: {most_expensive['Имя']} — цена {most_expensive['Цена']}, количество {most_expensive['Количество']}, стоимость {most_expensive['Стоимость']}"
    )

    doc.save('word.docx')

def convert_word_to_pdf():
    docx2pdf_convert('word.docx', 'pdf_file.pdf')

df = create_excel()
create_word(df)
convert_word_to_pdf()